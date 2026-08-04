"""
Plain text extraction from uploaded resume files. Kept dependency-light
(pypdf + python-docx) - no OCR, so scanned/image-only PDFs will extract
poorly; that's a known limitation, not a bug.

Works directly from bytes (via io.BytesIO) rather than a file path - no
temp files needed, and it's what lets files be stored in Mongo instead
of local disk (see resume_service.py for why: Render's free tier wipes
the filesystem on every deploy, which was silently breaking resume/
cover-letter attachments in application emails).
"""
import io
from pathlib import Path
from pypdf import PdfReader
from docx import Document


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(content)
    if ext == ".docx":
        return _extract_docx(content)
    raise ValueError(f"Unsupported resume file type: {ext}")


def _extract_pdf(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages).strip()


def _extract_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text]
    # also pull text out of tables (many resumes use table layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()
